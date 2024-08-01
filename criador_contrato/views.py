from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from .forms import ContratoForm
from docx import Document
from docx.shared import Pt, RGBColor
import os
import brazilcep
from datetime import datetime

def index(request):
    form = ContratoForm()
    return render(request, 'index.html', {'form': form})

def gerar_contrato(request):
    if request.method == 'POST':
        print("Recebida requisição POST")
        form = ContratoForm(request.POST)
        if form.is_valid():
            dados = form.cleaned_data
            print("Dados recebidos do formulário:", dados)  # Log para depuração

            # Monta o endereço completo
            endereco_completo = f"{dados['rua']}, {dados['cidade']}, {dados['estado']}, CEP: {dados['cep']}"
            dados['endereco'] = endereco_completo

            # Ajusta os dados conforme o gênero do administrador
            if dados['genero'] == 'Homem':
                dados['generoadm'] = 'representado pelo seu administrador'
                dados['inscrito'] = 'portador do'
            else:
                dados['generoadm'] = 'representada pela sua administradora'
                dados['inscrito'] = 'portadora do'

            # Adiciona cidade e data atual
            dados['cidade'] = dados.get('cidade', '')
            dados['data_atual'] = datetime.now().strftime('%d/%m/%Y')

            # Carrega o documento do modelo de contrato
            doc_path = os.path.join(os.path.dirname(__file__), 'modelo_contrato.docx')
            if not os.path.exists(doc_path):
                return HttpResponse("Modelo de contrato não encontrado.", status=404)
            
            try:
                doc = Document(doc_path)

                # Define o estilo para o texto das variáveis
                for paragrafo in doc.paragraphs:
                    for chave, valor in dados.items():
                        placeholder = f'{{{{ {chave} }}}}'
                        if placeholder in paragrafo.text:
                            for run in paragrafo.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, valor)
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Cor preta

                # Define o estilo para o texto "CONTRATANTE:"
                for paragrafo in doc.paragraphs:
                    if 'CONTRATANTE:' in paragrafo.text:
                        for run in paragrafo.runs:
                            if 'CONTRATANTE:' in run.text:
                                run.bold = True

                # Prepara a resposta com o documento gerado
                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=contrato_gerado.docx'
                
                # Salva o documento no buffer de resposta
                doc.save(response)
                
                return response

            except Exception as e:
                return HttpResponse(f"Erro ao gerar o documento: {str(e)}", status=500)

    return HttpResponse("Método não permitido.", status=405)

def consultar_cep(request):
    cep = request.GET.get('cep')
    if cep:
        try:
            endereco = brazilcep.get_address_from_cep(cep)
            return JsonResponse({
                'rua': endereco.get('street', ''),
                'cidade': endereco.get('city', ''),
                'estado': endereco.get('uf', '')
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'CEP não fornecido'}, status=400)
